#!/usr/bin/env python3
"""从支持的说明书格式中提取文本/表格，输出为中性 JSON 文档。

本模块是技能流水线第 1 步"解析说明书"的实现，负责把用户提供的设备说明书
（PDF / DOC / DOCX / XLS / XLSX）转换为统一的中性 JSON 结构，供后续
LLM 分析生成 points.json。

输出 JSON 结构：
  {
    "source": "<原始文件路径>",
    "format": "<文件格式后缀，如 pdf/docx/xlsx>",
    "text":   "<纯文本内容，Excel 为空字符串>",
    "tables": [<表格数据，结构因格式而异>]
  }

格式与依赖对应关系：
  - .docx  -> python-docx（段落 + 表格）
  - .xlsx  -> openpyxl（逐 sheet 读取行）
  - .xls   -> xlrd>=2.0.1（openpyxl 不支持旧版 .xls）
  - .doc   -> antiword / catdoc / LibreOffice（系统命令，非 python 库）
  - .pdf   -> PyMuPDF（fitz），扫描件需额外 OCR（本模块不含 OCR）

被调用：
  - 命令行：python scripts/extract_input.py <说明书路径> --output <结果.json>
  - 也可被其他 Python 脚本 import extract 后调用 extract(path) 函数
"""
from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from workspace import resolve_out


class InputFormatError(RuntimeError):
    """输入格式不支持或依赖缺失时抛出。

    调用方（main）会捕获此异常并通过 argparse 的 error() 以非零退出码终止，
    同时把错误信息打印给用户。
    """
    pass


def extract_doc(path):
    """提取旧版 .doc 文档文本（不含表格结构）。

    输入：
      path —— .doc 文件路径
    输出：
      {"text": "<文档全文>", "tables": []}
    作用：
      按优先级依次尝试系统命令 antiword -> catdoc -> LibreOffice(soffice)。
      python-docx 不支持 .doc，因此必须依赖外部工具。
      antiword/catdoc 输出字节流，尝试 utf-8 和 gb18030 两种编码解码。
      LibreOffice 以 headless 模式将 .doc 转为 txt 再读取。
    调用情况：
      被 extract() 在 suffix==".doc" 时调用；所有工具都不可用时抛 InputFormatError。
    """
    for executable, command in (
        ("antiword", lambda exe: [exe, str(path)]),
        ("catdoc", lambda exe: [exe, str(path)]),
    ):
        exe = shutil.which(executable)
        if exe:
            result = subprocess.run(command(exe), capture_output=True, check=True)
            for encoding in ("utf-8", "gb18030"):
                try:
                    return {"text": result.stdout.decode(encoding), "tables": []}
                except UnicodeDecodeError:
                    pass
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if office:
        with tempfile.TemporaryDirectory() as temp_dir:
            subprocess.run(
                [office, "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, str(path)],
                capture_output=True, check=True,
            )
            converted = Path(temp_dir) / (Path(path).stem + ".txt")
            if converted.exists():
                return {"text": converted.read_text(encoding="utf-8", errors="replace"), "tables": []}
    raise InputFormatError(
        "旧版 .doc 需要系统安装 antiword、catdoc 或 LibreOffice；python-docx 仅支持 .docx。"
    )


def extract_docx(path):
    """提取 .docx 文档的段落文本和表格。

    输入：
      path —— .docx 文件路径
    输出：
      {"text": "<段落拼接文本>", "tables": [<三维数组：表->行->单元格>]}
    作用：
      使用 python-docx 读取文档：
        - text：所有非空段落用换行符拼接
        - tables：每个表格转为 [[单元格文本, ...], ...] 的二维数组，多表组成三维数组
    调用情况：被 extract() 在 suffix==".docx" 时调用。
    """
    from docx import Document
    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs if p.text)
    tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
    return {"text": text, "tables": tables}


def extract_excel(path, legacy=False):
    """提取 Excel 工作簿所有 sheet 的行数据。

    输入：
      path   —— Excel 文件路径
      legacy —— 是否为旧版 .xls（True 用 xlrd，False 用 openpyxl）
    输出：
      {"text": "", "tables": [{"name": "<sheet名>", "rows": [[单元格值, ...], ...]}, ...]}
    作用：
      - .xlsx（legacy=False）：openpyxl 以 read_only 模式逐 sheet 读取，data_only 取公式计算值
      - .xls（legacy=True）：xlrd 读取，xlrd>=2.0.1 仅支持 .xls
      text 字段为空字符串（Excel 内容全部在 tables 中）
    调用情况：被 extract() 调用，根据后缀决定 legacy 参数。
    """
    sheets = []
    if legacy:
        try:
            import xlrd
        except ImportError as exc:
            raise InputFormatError("解析 .xls 需要安装 xlrd>=2.0.1。") from exc
        workbook = xlrd.open_workbook(path)
        for sheet in workbook.sheets():
            sheets.append({"name": sheet.name, "rows": [sheet.row_values(i) for i in range(sheet.nrows)]})
    else:
        import openpyxl
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
        for sheet in workbook.worksheets:
            sheets.append({"name": sheet.title, "rows": [list(row) for row in sheet.iter_rows(values_only=True)]})
        workbook.close()
    return {"text": "", "tables": sheets}


def extract_pdf(path):
    """提取 PDF 文档的文本内容（不含表格结构）。

    输入：
      path —— PDF 文件路径
    输出：
      {"text": "<全文文本，各页用换行拼接>", "tables": []}
    作用：
      使用 PyMuPDF（fitz）逐页提取文本；扫描件 PDF 无文本层时返回空文本，
      需要调用方额外做 OCR（本模块不含 OCR 能力）。
    调用情况：被 extract() 在 suffix==".pdf" 时调用。
    """
    import fitz
    document = fitz.open(path)
    try:
        return {"text": "\n".join(page.get_text() for page in document), "tables": []}
    finally:
        document.close()


def extract(path):
    """根据文件后缀分派到对应的提取器，返回中性 JSON 结构。

    输入：
      path —— 说明书文件路径（支持 .pdf/.doc/.docx/.xls/.xlsx）
    输出：
      {"source": "<原始路径>", "format": "<后缀无点>", "text": "...", "tables": [...]}
    作用：
      - 校验文件存在性
      - 按后缀选择 handler（.xls 走 legacy=True 的 extract_excel）
      - 合并 source 和 format 元信息后返回
    调用情况：
      - 被 main() 调用
      - 可被其他脚本 import 后直接调用
    异常：文件不存在或不支持的后缀抛 InputFormatError。
    """
    path = Path(path)
    if not path.is_file():
        raise InputFormatError(f"输入文件不存在: {path}")
    suffix = path.suffix.lower()
    handlers = {
        ".doc": extract_doc,
        ".docx": extract_docx,
        ".xls": lambda p: extract_excel(p, legacy=True),
        ".xlsx": extract_excel,
        ".pdf": extract_pdf,
    }
    if suffix not in handlers:
        raise InputFormatError(f"不支持的输入格式: {suffix}; 支持 .pdf/.doc/.docx/.xls/.xlsx")
    result = handlers[suffix](path)
    return {"source": str(path), "format": suffix[1:], **result}


def main():
    """命令行入口：解析参数，调用 extract，输出 JSON。

    用法：
      python extract_input.py <说明书路径> [--output <结果.json>]
    行为：
      - --output 缺省时打印 JSON 到 stdout
      - --output 指定时写入文件（UTF-8，缩进 2）
      - 提取失败（格式不支持/依赖缺失/命令执行失败）时以非零退出码终止
    """
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("--output", help="可选 JSON 输出路径；缺省打印到 stdout")
    args = parser.parse_args()
    try:
        result = extract(args.input)
    except (InputFormatError, subprocess.CalledProcessError) as exc:
        parser.error(str(exc))
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    out_path = args.output or resolve_out(basename="extract_output.json")
    Path(out_path).write_text(payload + "\n", encoding="utf-8")
    print("saved:", out_path)


if __name__ == "__main__":
    main()
